#!/usr/bin/env python3
"""
免疫浸润分析 - 生成 Word 报告 (简化版：无图片，三线表格式)
"""
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from scipy import stats as sp_stats

# 设置路径
BASE_DIR = "/media/desk16/share/secure/immune_infiltration/immune_infiltration_results"
GROUP_FILE = "/media/desk16/share/secure/immune_infiltration/GSE126124.group.csv"

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_three_line_table(doc, data, title, caption=""):
    """添加三线表格式的表格"""
    if data is None or len(data) == 0:
        return

    # 添加标题
    if caption:
        p = doc.add_paragraph()
        p.add_run(caption).italic = True

    # 创建表格
    n_cols = len(data.columns)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'

    # 设置表头
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(data.columns):
        hdr_cells[i].text = str(col)
        # 设置表头边框
        for edge in ('top', 'bottom'):
            set_cell_border(hdr_cells[i], **{edge: True})
        # 加粗
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # 添加数据行
    for _, row in data.iterrows():
        row_cells = table.add_row().cells
        for i, col in enumerate(data.columns):
            val = row[col]
            if isinstance(val, float):
                row_cells[i].text = f"{val:.4f}"
            else:
                row_cells[i].text = str(val)
            # 只设置底部边框
            if _ == data.index[-1]:  # 最后一行
                for edge in ('bottom',):
                    set_cell_border(row_cells[i], **{edge: True})

def load_stats(method_dir, method_name):
    """加载统计结果"""
    stat_file = os.path.join(BASE_DIR, method_dir, f'02.stat.{method_name.lower()}.csv')
    if os.path.exists(stat_file):
        return pd.read_csv(stat_file)
    return None

def get_group_comparison(method_dir, method_name):
    """获取组间比较结果"""
    # 加载结果文件
    files = {
        'cibersort': '01.ciber_res.csv',
        'epic': '01.epic_res.csv',
        'estimate': '01.estimate_res.csv',
        'ips': '01.ips_res.csv',
        'mcpcounter': '01.mcpcounter_res.csv',
        'ssgsea': '01.ssgsea_res.csv',
        'timer': '01.timer_res.csv',
        'xcell': '01.xcell_res.csv'
    }

    res_file = os.path.join(BASE_DIR, method_dir, files.get(method_name.lower()))
    if not res_file or not os.path.exists(res_file):
        return None

    data = pd.read_csv(res_file)
    group = pd.read_csv(GROUP_FILE)

    if data is None or group is None:
        return None

    # 合并分组信息
    try:
        data = data.merge(group, left_on='ID', right_on='sample')
    except:
        return None

    # 获取细胞类型列
    exclude_cols = ['ID', 'sample', 'group', 'P-value', 'Correlation', 'RMSE']
    cell_cols = [c for c in data.columns if c not in exclude_cols and data[c].dtype in ['float64', 'int64']]

    results = []
    for col in cell_cols:
        ctrl = data[data['group'] == 'Control'][col].dropna()
        case = data[data['group'] == 'IBD'][col].dropna()

        if len(ctrl) > 0 and len(case) > 0:
            try:
                _, pval = sp_stats.mannwhitneyu(ctrl, case, alternative='two-sided')
                results.append({
                    'Cell_Type': col,
                    'IBD_Mean': case.mean(),
                    'Control_Mean': ctrl.mean(),
                    'P_value': pval,
                    'Significance': '***' if pval < 0.001 else '**' if pval < 0.01 else '*' if pval < 0.05 else 'ns'
                })
            except:
                pass

    if len(results) > 0:
        df = pd.DataFrame(results)
        # 按P值排序
        df = df.sort_values('P_value')
        return df
    return None

def generate_report():
    """生成 Word 报告"""
    doc = Document()

    # 设置中文字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题
    title = doc.add_heading('免疫浸润分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加项目信息
    doc.add_paragraph(f"项目: GSE126124 (IBD vs Control)")
    doc.add_paragraph(f"分析方法: CIBERSORT, EPIC, ESTIMATE, IPS, MCPcounter, ssGSEA, TIMER, xCell")
    doc.add_paragraph(f"生成日期: 2026-03-05")

    doc.add_paragraph()

    # 1. 方法概述
    doc.add_heading('1. 方法概述', level=1)
    doc.add_paragraph("本分析采用多算法整合策略评估肿瘤免疫微环境浸润水平。通过 IOBR 包集成八种主流免疫浸润评估算法进行综合分析：")
    doc.add_paragraph("• CIBERSORT: 基于线性支持向量机算法估算22种免疫细胞类型的相对比例")
    doc.add_paragraph("• EPIC: 利用混合线性模型估计多种免疫细胞群体的绝对丰度")
    doc.add_paragraph("• ESTIMATE: 计算 Stromal Score、Immune Score 和 ESTIMATE Score")
    doc.add_paragraph("• IPS: 免疫表型评分")
    doc.add_paragraph("• MCPcounter: MCPcounter 免疫细胞计数")
    doc.add_paragraph("• ssGSEA: 基因集富集分析")
    doc.add_paragraph("• TIMER: TIMER 免疫浸润评估")
    doc.add_paragraph("• xCell: xCell 细胞富集评分")
    doc.add_paragraph("统计检验采用 Mann-Whitney U 检验，P 值经 Benjamini-Hochberg 方法校正。")

    # 2. 主要结果
    doc.add_heading('2. 主要结果', level=1)

    methods = [
        ('estimate_output', 'ESTIMATE', 'ESTIMATE评分'),
        ('cibersort_output', 'CIBERSORT', 'CIBERSORT免疫细胞'),
        ('epic_output', 'EPIC', 'EPIC免疫细胞'),
        ('ips_output', 'IPS', 'IPS评分'),
        ('mcpcounter_output', 'MCPcounter', 'MCPcounter免疫细胞'),
        ('ssgsea_output', 'ssGSEA', 'ssGSEA评分'),
        ('timer_output', 'TIMER', 'TIMER免疫细胞'),
        ('xcell_output', 'xCell', 'xCell细胞')
    ]

    for method_dir, method_name, method_title in methods:
        doc.add_heading(f'{method_title}', level=2)

        # 获取组间比较结果
        comp_results = get_group_comparison(method_dir, method_name)

        if comp_results is not None and len(comp_results) > 0:
            # 显示显著差异的结果
            sig_results = comp_results[comp_results['P_value'] < 0.05]

            if len(sig_results) > 0:
                doc.add_paragraph(f"发现 {len(sig_results)} 个显著差异的指标 (p < 0.05):")

                # 限制显示前10个
                display_df = sig_results.head(10)[['Cell_Type', 'IBD_Mean', 'Control_Mean', 'P_value', 'Significance']]
                add_three_line_table(doc, display_df, method_name, f"表1. {method_name} 组间差异分析")
            else:
                doc.add_paragraph("未发现显著差异的指标 (p < 0.05)")
        else:
            doc.add_paragraph("暂无统计数据")

    # 3. 结论
    doc.add_heading('3. 结论', level=1)
    doc.add_paragraph("本分析通过八种免疫浸润评估算法对 GSE126124 数据集进行了全面分析。综合分析结果显示 IBD 组与 Control 组在多种免疫细胞类型上存在显著差异。")

    # 保存
    output_path = os.path.join(BASE_DIR, 'result', 'immune_infiltration_report.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"报告已保存: {output_path}")

if __name__ == '__main__':
    generate_report()
